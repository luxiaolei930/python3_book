#导入docx库
from docx import Document
from docx.shared import Inches
#文档路径为document.docx
path = "document.docx"
#打开document.docx
document = Document(path)
#读取文档中所有段落文本
for p in document.paragraphs:
    #打印文档中所有段落文本
    print(p.text)
#获取文件中的表格
tables = document.tables 
#获取文件中的第一个表格（如需读取文档中所有表格，可使用for循环遍历tables）
table = tables[0]
#从表格第一行开始循环读取表格，len(table.rows)表示table的行数。使用for循环遍历第1行到第n行，n为表格行数。
for i in range(0,len(table.rows)):
    #cell为单元格，cell(i,0)表示第(i+1)行第1列单元格，以此类推；+表示字符串的拼接，将三个单元格中的内容拼接；" "表示使用空格隔开；
    result = table.cell(i,0).text + " " +table.cell(i,1).text+" "+table.cell(i,2).text
    print(result)
