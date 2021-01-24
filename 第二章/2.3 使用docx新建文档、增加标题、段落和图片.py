#导入docx库
from docx import Document
from docx.shared import Inches
#新建Document对象
document = Document()
#增加文档标题
document.add_heading('文档标题：Python 3 语料库技术与应用', 0)
#增加一个段落
p = document.add_paragraph('添加段落：DOC和DOCX是Microsoft Office word的两种文档格式。Python提供了一些现成的库可以处理word文档。本文采用python-docx，如果需要doc文档，可先将docx先转换成doc。使用python批量处理docx文档，可实现新建Document对象、增加文档标题、增加段落、添加图片、添加表格等功能 ')
p.add_run('在这里添加一句加粗文字').bold = True #添加加粗文字
p.add_run('在这里添加一句斜体文字.').italic = True #添加斜体文字
#增加文档一级标题
document.add_heading('一级标题', level=1)
#增加一个段落 
document.add_paragraph('这个段落的格式为：斜体+下划线。DOC和DOCX是Microsoft Office word的两种文档格式。Python提供了一些现成的库可以处理word文档。本文采用python-docx，如果需要doc文档，可先将docx先转换成doc。使用python批量处理docx文档，可实现新建Document对象、增加文档标题、增加段落、添加图片、添加表格等功能', style='Intense Quote')
#增加一个段落-列表形式
document.add_paragraph(
    '添加无数字编号的文字列表', style='List Bullet'
)
document.add_paragraph(
    '添加有数字编号的文字列表', style='List Number'
)
#添加图片
document.add_picture('qianqian.png', width=Inches(2.25))
#保存文档
document.save('document.docx')