#导入docx库  
from docx import Document  
from docx.shared import Inches  
  
#新建Document对象  
document = Document()  
	  
#增加文档标题  
document.add_heading('添加表格', 3)  
	  
#添加表格  
table_content = (  
	('001', '小狗浅浅',1 ),  
	('002', '玉米',7),  
	('003', '红薯',2)  
)  
	  
table = document.add_table(rows=1, cols=3)  
header_cells = table.rows[0].cells  
header_cells[0].text = '编号'  
header_cells[1].text = '内容'  
header_cells[2].text = '数量'  
for number, content, quantity in table_content:  
	row_cells = table.add_row().cells  
	row_cells[0].text = str(number)  
	row_cells[1].text = content  
	row_cells[2].text = str(quantity)  
#保存文档  
document.save('document2.docx')     
